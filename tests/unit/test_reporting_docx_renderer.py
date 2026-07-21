"""Unit tests for core/reporting/docx_renderer.py."""

from __future__ import annotations

import base64
import io

import plotly.graph_objects as go
import pytest
from docx import Document

from core.reporting.docx_renderer import DOCXReportRenderer
from core.reporting.exceptions import ChartRenderingError
from core.reporting.models import (
    GeneratedReport,
    ReportSection,
    ReportSectionType,
    ReportStatistics,
    ReportType,
    ReportValidationResult,
)

pytestmark = pytest.mark.unit

_TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


class _FakeChartImageEncoder:
    def encode(self, figure: go.Figure, *, width: int = 900, height: int = 500) -> bytes:
        return _TINY_PNG


class _FailingChartImageEncoder:
    def encode(self, figure: go.Figure, *, width: int = 900, height: int = 500) -> bytes:
        raise ChartRenderingError("no chrome available")


def _report() -> GeneratedReport:
    sections = (
        ReportSection(
            section_type=ReportSectionType.FINDINGS,
            title="Findings",
            content={
                "finding_count": 1,
                "findings": [{"title": "Brute force", "severity": "high", "risk_score": 70.0}],
            },
        ),
    )
    return GeneratedReport(
        case_id="case-1",
        report_type=ReportType.TECHNICAL_INVESTIGATION,
        title="Technical Investigation Report",
        sections=sections,
        statistics=ReportStatistics(finding_count=1),
        validation=ReportValidationResult(is_complete=True),
        confidence=0.6,
    )


def test_render_produces_a_valid_docx_document() -> None:
    renderer = DOCXReportRenderer(chart_image_encoder=_FakeChartImageEncoder())
    docx_bytes = renderer.render(_report(), include_charts=False)
    document = Document(io.BytesIO(docx_bytes))
    headings = [
        p.text
        for p in document.paragraphs
        if p.style.name.startswith("Heading") or p.style.name == "Title"
    ]
    assert any("Technical Investigation Report" in h for h in headings)
    assert any("Findings" in h for h in headings)


def test_render_includes_toc_field() -> None:
    renderer = DOCXReportRenderer(chart_image_encoder=_FakeChartImageEncoder())
    docx_bytes = renderer.render(_report(), include_charts=False)
    document = Document(io.BytesIO(docx_bytes))
    xml = document.element.xml
    assert "TOC" in xml


def test_render_includes_a_findings_table() -> None:
    renderer = DOCXReportRenderer(chart_image_encoder=_FakeChartImageEncoder())
    docx_bytes = renderer.render(_report(), include_charts=False)
    document = Document(io.BytesIO(docx_bytes))
    assert len(document.tables) >= 1
    header_row_text = [cell.text for cell in document.tables[0].rows[0].cells]
    assert "title" in header_row_text


def test_render_degrades_gracefully_when_chart_rendering_fails() -> None:
    renderer = DOCXReportRenderer(chart_image_encoder=_FailingChartImageEncoder())
    docx_bytes = renderer.render(_report(), include_charts=True)
    document = Document(io.BytesIO(docx_bytes))  # never crashes, still a valid docx
    assert document is not None


def test_render_embeds_organization_branding() -> None:
    from core.reporting.theme import LIGHT_THEME

    theme = LIGHT_THEME.model_copy(
        update={"organization_name": "Acme Corp", "footer_text": "Confidential"}
    )
    renderer = DOCXReportRenderer(chart_image_encoder=_FakeChartImageEncoder())
    docx_bytes = renderer.render(_report(), theme=theme, include_charts=False)
    document = Document(io.BytesIO(docx_bytes))
    all_text = "\n".join(p.text for p in document.paragraphs)
    assert "Acme Corp" in all_text
    assert "Confidential" in all_text

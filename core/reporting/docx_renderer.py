"""DOCX Renderer — the task's named "DOCX Renderer".

Built on `python-docx` (already a project dependency, used elsewhere for
parsing `.docx` incident notes — this is its first use for *generating* a
document). Produces a genuinely editable Word document: real heading
styles (so Word's own outline/navigation pane works), real tables, a
Table-of-Contents *field* Word recomputes on open (python-docx has no
first-class TOC API; the field-code approach below is the standard,
documented recipe for inserting one — see `_add_toc_field`'s docstring for
why it needs a raw XML fragment).
"""

from __future__ import annotations

import io

from docx import Document as _new_document  # noqa: N813
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement  # type: ignore[attr-defined]  # not in python-docx's stubs
from docx.shared import Inches, Pt, RGBColor

from core.reporting.asset_manager import AssetManager, from_data_uri
from core.reporting.chart_image_encoder import (
    ChartImageEncoder,
    KaleidoChartImageEncoder,
    safe_encode,
)
from core.reporting.charts import build_all_charts
from core.reporting.models import GeneratedReport, ReportSection
from core.reporting.theme import ReportTheme, resolve_theme


def _hex_to_rgbcolor(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def _add_toc_field(document: Document) -> None:
    """Inserts a `TOC \\o "1-3" \\h \\z \\u` field — Word populates this the
    first time the document is opened and the user accepts the "update
    field" prompt (or presses F9). python-docx has no `add_toc()` API
    because a TOC is a *computed* Word field, not static content this
    library can pre-render without duplicating Word's own pagination
    engine; inserting the raw field code and letting Word compute page
    numbers itself is the correct, standard approach, not a workaround for
    a bug."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' to generate the table of contents."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def _render_value(document: Document, value: object) -> None:
    if isinstance(value, dict):
        for key, item in value.items():
            if isinstance(item, dict | list | tuple):
                document.add_paragraph(str(key), style="Intense Quote")
                _render_value(document, item)
            else:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(f"{key}: ").bold = True
                paragraph.add_run(str(item))
        return
    if isinstance(value, list | tuple):
        if not value:
            document.add_paragraph("(none)")
            return
        if all(isinstance(item, dict) for item in value):
            _add_table(document, list(value))
            return
        for item in value:
            document.add_paragraph(str(item), style="List Bullet")
        return
    document.add_paragraph(str(value))


def _add_table(document: Document, rows: list[dict[str, object]]) -> None:
    headers = list(rows[0].keys())
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for index, header in enumerate(headers):
            cells[index].text = str(row.get(header, ""))


def _render_section(document: Document, section: ReportSection) -> None:
    document.add_heading(section.title, level=2)
    if section.is_empty:
        paragraph = document.add_paragraph()
        run = paragraph.add_run("No data available for this section.")
        run.italic = True
        return
    _render_value(document, section.content)


class DOCXReportRenderer:
    def __init__(
        self,
        *,
        asset_manager: AssetManager | None = None,
        chart_image_encoder: ChartImageEncoder | None = None,
    ) -> None:
        self._asset_manager = asset_manager or AssetManager()
        self._chart_image_encoder = chart_image_encoder or KaleidoChartImageEncoder()

    def render(
        self,
        report: GeneratedReport,
        *,
        theme: ReportTheme | str | None = None,
        include_charts: bool = True,
    ) -> bytes:
        resolved_theme = resolve_theme(theme)
        document: Document = _new_document()

        title_style = document.styles["Title"]
        title_style.font.color.rgb = _hex_to_rgbcolor(resolved_theme.primary_color)

        if resolved_theme.logo_data_uri:
            logo_bytes = from_data_uri(resolved_theme.logo_data_uri)
            if logo_bytes is not None:
                document.add_picture(io.BytesIO(logo_bytes), width=Inches(1.5))
        if resolved_theme.organization_name:
            document.add_paragraph(resolved_theme.organization_name)

        document.add_heading(report.title, level=0)
        meta = document.add_paragraph()
        meta.add_run(
            f"Case: {report.case_id} | Type: {report.report_type.value} | "
            f"Generated: {report.generated_at.isoformat()}"
        ).italic = True
        status = "Degraded" if report.degraded else "Complete"
        status_paragraph = document.add_paragraph()
        status_paragraph.add_run(f"Confidence: {report.confidence * 100:.0f}% | Status: {status}")

        document.add_heading("Table of Contents", level=1)
        _add_toc_field(document)
        document.add_page_break()  # type: ignore[no-untyped-call]  # python-docx stub gap

        for section in report.sections:
            _render_section(document, section)

        if include_charts:
            document.add_page_break()  # type: ignore[no-untyped-call]  # python-docx stub gap
            document.add_heading("Charts", level=1)
            for name, figure in build_all_charts(report).items():
                title = name.replace("_", " ").title()
                document.add_heading(title, level=2)
                image_bytes = safe_encode(self._chart_image_encoder, figure, chart_name=name)
                if image_bytes is None:
                    paragraph = document.add_paragraph()
                    paragraph.add_run("Chart rendering unavailable.").italic = True
                    continue
                validated = self._asset_manager.prepare_binary(image_bytes, mime_type="image/png")
                document.add_picture(io.BytesIO(validated), width=Inches(6.0))

        document.add_heading("Appendix", level=1)
        footer_paragraph = document.add_paragraph()
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_paragraph.add_run(f"Report ID: {report.report_id}")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = _hex_to_rgbcolor(resolved_theme.secondary_color)
        if resolved_theme.footer_text:
            document.add_paragraph(resolved_theme.footer_text)

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()
